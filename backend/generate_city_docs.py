"""
Generate licensing/compliance .docx documents for multiple Indian cities.
Each document contains city-specific restaurant licensing rules, building regulations,
fire safety norms, FSSAI requirements, and zoning rules.
Run once: python generate_city_docs.py
"""

import os
from docx import Document

DATA_DIR = os.path.join(os.path.dirname(__file__), "data")
os.makedirs(DATA_DIR, exist_ok=True)

CITY_DOCS = {
    "Mumbai": {
        "filename": "Mumbai Restaurant Licensing & Compliance Guide.docx",
        "content": [
            ("Mumbai Restaurant Licensing & Compliance Guide (Official Sources)", True),
            ("Overview", True),
            ("This document outlines the regulatory requirements for opening and operating a restaurant or food business in Mumbai, Maharashtra. "
             "All businesses must comply with BMC (Brihanmumbai Municipal Corporation), Maharashtra state, and central government regulations.", False),

            ("1. FSSAI License / Registration", True),
            ("All food businesses in Mumbai require FSSAI compliance. "
             "Turnover up to ₹12 lakh per year: FSSAI Registration (Form A). "
             "Turnover above ₹12 lakh per year: FSSAI State License (Form B). "
             "Turnover above ₹20 crore: FSSAI Central License. "
             "Application is online via the FoSCoS portal. "
             "Documents required: ID proof, address proof, food safety management plan, NOC from local authority.", False),

            ("2. BMC Shop & Establishment License", True),
            ("Every restaurant in Mumbai must obtain a Shop & Establishment License from BMC within 30 days of opening. "
             "Apply through the BMC portal or ward office. "
             "Required: PAN card, address proof, partnership deed/MOA, photos of premises. "
             "Annual renewal is mandatory. Late renewal attracts a penalty.", False),

            ("3. BMC Health/Trade License", True),
            ("A Health Trade License from BMC is mandatory for all food establishments. "
             "Inspection by BMC health officer is required before issuance. "
             "The premises must meet hygiene standards: clean kitchen, pest control measures, proper waste disposal. "
             "Water quality certificate from BMC is required. "
             "License must be renewed annually.", False),

            ("4. Fire NOC (No Objection Certificate)", True),
            ("Fire NOC from Mumbai Fire Brigade is required for: "
             "All restaurants with built-up area above 500 sq ft. "
             "All restaurants above ground floor or in multi-story buildings. "
             "Restaurants with seating capacity exceeding 50 persons. "
             "Requirements: Minimum 2 exits for restaurants above 1000 sq ft. "
             "Fire extinguishers (ABC type) — 1 per 1000 sq ft or part thereof. "
             "Emergency exit signage with illumination. "
             "Fire-resistant kitchen hood and suppression system for commercial kitchens. "
             "Staircase width minimum 1.2 meters for buildings above ground floor. "
             "Annual fire safety audit is mandatory for establishments above 2000 sq ft.", False),

            ("5. Building Plan Approval", True),
            ("Any new construction or structural alteration requires approved building plans from BMC. "
             "For commercial restaurants, the plan must be approved by a licensed architect registered with the Council of Architecture. "
             "Buildings in CRZ (Coastal Regulation Zone) areas require additional clearance from MCZMA (Maharashtra Coastal Zone Management Authority). "
             "FSI (Floor Space Index) for commercial use in Mumbai: varies by zone — typically 1.33 to 5.0 depending on location. "
             "Maximum building height depends on road width and zoning (DCR 2034 regulations). "
             "Minimum floor-to-ceiling height for commercial restaurants: 3.0 meters (ground floor), 2.75 meters (upper floors). "
             "Minimum kitchen area: 15% of total dining area or 15 sq meters, whichever is greater.", False),

            ("6. Liquor License", True),
            ("Liquor license in Mumbai is issued by the Maharashtra State Excise Department. "
             "FL-III license for serving Indian-made liquor. "
             "FL-IV license for serving foreign liquor (IMFL). "
             "Restaurant must have minimum carpet area of 1000 sq ft for FL-III and 1500 sq ft for FL-IV. "
             "Cannot be located within 100 meters of religious places, educational institutions, or hospitals. "
             "Serving hours: 11:00 AM to 1:30 AM (may vary by area). "
             "Annual renewal required with fee payment.", False),

            ("7. Eating House License", True),
            ("An Eating House License from Mumbai Police (issued by Commissioner of Police) is mandatory for all restaurants in Mumbai. "
             "Police verification of the owner and premises is conducted. "
             "CCTV cameras are mandatory at entry/exit points. "
             "Application requires: identity proof, address proof, FSSAI license copy, BMC trade license copy, fire NOC.", False),

            ("8. Signage & Hoarding License", True),
            ("Any signage or hoarding requires BMC permission. "
             "Maximum signage area and illumination rules apply per BMC guidelines. "
             "Digital/LED signs require additional electrical safety certification.", False),

            ("9. GST Registration", True),
            ("GST registration is mandatory for restaurants with turnover exceeding ₹20 lakh (₹10 lakh for NE states). "
             "Restaurants not serving alcohol: 5% GST without ITC. "
             "Restaurants in hotels with room tariff above ₹7,500: 18% GST with ITC.", False),

            ("10. Environmental Clearance", True),
            ("Restaurants with significant noise or air emissions may require NOC from Maharashtra Pollution Control Board (MPCB). "
             "Kitchen exhaust must have grease traps and filters. "
             "Noise levels must not exceed 55 dB during daytime and 45 dB at nighttime at the boundary.", False),

            ("Key Compliance Checkpoints for Blueprint Review", True),
            ("Floor plan must clearly show: kitchen area, dining area, all exits, restrooms, fire equipment locations. "
             "Kitchen must be present and clearly demarcated for any food-serving establishment. "
             "Minimum 2 exits for area above 1000 sq ft. "
             "Emergency exit width: minimum 1.0 meter. "
             "Restroom facilities: minimum 1 per 50 seats (separate for men and women if seating > 30). "
             "The floor plan must be prepared by a licensed architect registered with the Council of Architecture.", False),
        ],
    },

    "Bangalore": {
        "filename": "Bangalore Restaurant Licensing & Compliance Guide.docx",
        "content": [
            ("Bangalore Restaurant Licensing & Compliance Guide (Official Sources)", True),
            ("Overview", True),
            ("This document covers regulatory requirements for restaurants and food businesses in Bangalore (Bengaluru), Karnataka. "
             "Businesses must comply with BBMP (Bruhat Bengaluru Mahanagara Palike), Karnataka state, and central government regulations.", False),

            ("1. FSSAI License / Registration", True),
            ("All food businesses in Bangalore must have FSSAI compliance. "
             "Turnover ≤ ₹12 lakh: FSSAI Registration. "
             "Turnover > ₹12 lakh: FSSAI State License. "
             "Apply online via FoSCoS portal. "
             "Food safety supervisor must be appointed for licensed establishments.", False),

            ("2. BBMP Trade License", True),
            ("Trade License from BBMP is mandatory. Apply through the BBMP Sakala portal online. "
             "Categories: small (< 500 sq ft), medium (500-2000 sq ft), large (> 2000 sq ft). "
             "Different fee slabs apply. Annual renewal required. "
             "Required documents: property ownership/lease deed, ID proof, photos, building plan approval.", False),

            ("3. BBMP Health License", True),
            ("Health License from BBMP Health Department is required for all food establishments. "
             "Premises inspection by BBMP Health Inspector. "
             "Requirements: proper drainage, clean water supply, pest control certificate, waste segregation compliance.", False),

            ("4. Fire Safety NOC", True),
            ("Fire NOC from Karnataka State Fire & Emergency Services (KSFES) is required for: "
             "All commercial establishments with area exceeding 500 sq ft. "
             "All buildings above 15 meters height. "
             "Restaurants with seating above 50 persons. "
             "Requirements: Fire extinguishers — 1 ABC type per 1000 sq ft. "
             "Two independent exits for area above 1000 sq ft. "
             "Illuminated emergency exit signage. "
             "Automatic fire alarm for buildings above 2 floors. "
             "Sprinkler system for buildings above 4 floors or area above 5000 sq ft. "
             "Annual fire safety certificate renewal. "
             "Staircase width: minimum 1.25 meters.", False),

            ("5. Building Plan Approval", True),
            ("Building plan approval from BBMP or BDA (Bangalore Development Authority) is required for new construction. "
             "Licensed architect must prepare the plan. "
             "Maximum FAR (Floor Area Ratio) for commercial zones: 1.75 to 3.25 depending on road width. "
             "Minimum floor-to-ceiling height: 3.0 meters for ground floor commercial, 2.75 meters for upper floors. "
             "Setback requirements vary by zone — minimum 3 meters front setback for buildings on roads < 12 meters. "
             "Revised Master Plan 2031 (RMP-2031) regulations apply.", False),

            ("6. Karnataka Excise License for Liquor", True),
            ("CL-9 license for serving liquor in restaurants (Karnataka Excise Act). "
             "Minimum carpet area of 1000 sq ft required. "
             "Cannot be within 100 meters of religious/educational institutions. "
             "Online application through Karnataka Excise Department portal.", False),

            ("7. Police Permission / Eating House License", True),
            ("Police permission required for operating a restaurant. "
             "Background verification of owner. "
             "CCTV at entry/exit points mandatory for restaurants operating after 11 PM.", False),

            ("8. Environmental Compliance", True),
            ("NOC from Karnataka State Pollution Control Board (KSPCB) for restaurants with significant emissions. "
             "Sewage treatment compliance required for establishments above 2000 sq ft. "
             "Kitchen exhaust with grease traps mandatory. "
             "Noise limit: 55 dB daytime, 45 dB nighttime at boundary.", False),

            ("9. GST Registration", True),
            ("Same central rules: mandatory above ₹20 lakh turnover. "
             "5% GST for standalone restaurants without ITC.", False),

            ("Key Compliance Checkpoints for Blueprint Review", True),
            ("Floor plan must show kitchen, dining, all exits, restrooms, ventilation systems. "
             "Kitchen must be clearly marked and meet FSSAI standards. "
             "Minimum 2 exits for area above 1000 sq ft. "
             "Parking provision: 1 car space per 50 sq meters of floor area as per BBMP norms. "
             "Wheelchair accessibility required for new public buildings under RPWD Act 2016. "
             "Rainwater harvesting mandatory for all buildings on sites above 2400 sq ft.", False),
        ],
    },

    "Delhi": {
        "filename": "Delhi Restaurant Licensing & Compliance Guide.docx",
        "content": [
            ("Delhi Restaurant Licensing & Compliance Guide (Official Sources)", True),
            ("Overview", True),
            ("This document covers regulations for restaurants in Delhi NCR. "
             "Businesses must comply with MCD (Municipal Corporation of Delhi) / NDMC, Delhi state, and central government rules. "
             "Delhi follows the unified MCD structure after 2022 merger.", False),

            ("1. FSSAI License / Registration", True),
            ("Mandatory for all food businesses. "
             "Turnover up to ₹12 lakh: Registration. Turnover > ₹12 lakh: State License. "
             "Central License for turnover above ₹20 crore or multi-state operations. "
             "Online application via FoSCoS. "
             "Food safety supervisor training certificate required for licensed businesses.", False),

            ("2. MCD Health Trade License", True),
            ("Health Trade License from MCD is mandatory. "
             "Apply online through MCD portal. "
             "Inspection by Health Inspector before issuance. "
             "Hygiene requirements: clean kitchen, proper ventilation, pest control, waste management. "
             "Structural stability certificate for buildings older than 15 years. "
             "Annual renewal with compliance certificate.", False),

            ("3. Delhi Police Eating House License", True),
            ("Eating House License from Delhi Police (Licensing Unit) is mandatory. "
             "Required for all restaurants, cafes, dhabas serving food to public. "
             "Police verification of owner and employees. "
             "CCTV cameras mandatory. "
             "License valid for 1 year, renewable.", False),

            ("4. Fire NOC", True),
            ("Fire NOC from Delhi Fire Services is mandatory for: "
             "All restaurants with area exceeding 500 sq ft. "
             "All buildings above ground floor level. "
             "Seating capacity above 50 persons. "
             "Requirements: Minimum 2 independent exits for area above 500 sq ft. "
             "Exit width minimum 1.0 meter. "
             "Fire extinguishers: 1 per 500 sq ft for kitchen, 1 per 1000 sq ft for dining. "
             "Emergency lighting at all exits. "
             "Fire-resistant materials in kitchen area. "
             "Staircase width: minimum 1.5 meters for public buildings. "
             "Annual fire audit mandatory. "
             "Basement restaurants require special clearance with mechanical ventilation.", False),

            ("5. Building Plan and Zoning", True),
            ("Building plan sanction from MCD/DDA required for any new construction or change of use. "
             "Delhi Master Plan (MPD-2041) zoning regulations apply. "
             "Maximum FAR: 350 for commercial, 200 for residential zones. "
             "Minimum floor height: 2.75 meters. "
             "Restaurant is allowed in: Commercial zones, Mixed-use zones, and designated commercial streets in residential areas. "
             "Not allowed in: Purely residential zones (unless the area has been redesignated). "
             "Parking: 1 ECS (Equivalent Car Space) per 50 sq meters of floor area. "
             "All commercial buildings must have approved structural stability certificate.", False),

            ("6. Liquor License", True),
            ("L-17 license for restaurants serving liquor (Delhi Excise Department). "
             "Minimum covered area: 1000 sq ft. "
             "Cannot be within 100 meters of religious places, schools, or hospitals. "
             "Dry days compliance mandatory. "
             "Online application through Delhi Excise portal.", False),

            ("7. Pollution & Environment", True),
            ("Consent from Delhi Pollution Control Committee (DPCC) for restaurants with tandoor/charcoal cooking. "
             "Kitchen exhaust: chimney height must be above rooftop level. "
             "Grease traps mandatory. "
             "Noise limit: 55 dB daytime, 45 dB night at boundary. "
             "DG set (if any) must comply with emission norms.", False),

            ("8. GST Registration", True),
            ("Mandatory above ₹20 lakh turnover. 5% GST (no ITC) for standalone restaurants.", False),

            ("Key Compliance Checkpoints for Blueprint Review", True),
            ("Floor plan must clearly show: kitchen, store, dining, exits, restrooms, fire equipment. "
             "Kitchen must be present and adequately ventilated. "
             "Minimum 2 exits for area above 500 sq ft. "
             "Basement usage requires special fire clearance and mechanical ventilation. "
             "Parking provision as per MPD-2041. "
             "Structural stability certificate required for older buildings. "
             "Disabled-friendly access mandatory under RPWD Act.", False),
        ],
    },

    "Hyderabad": {
        "filename": "Hyderabad Restaurant Licensing & Compliance Guide.docx",
        "content": [
            ("Hyderabad Restaurant Licensing & Compliance Guide (Official Sources)", True),
            ("Overview", True),
            ("This document covers restaurant licensing regulations in Hyderabad, Telangana. "
             "Businesses must comply with GHMC (Greater Hyderabad Municipal Corporation), Telangana state, and central government rules.", False),

            ("1. FSSAI License / Registration", True),
            ("All food businesses in Hyderabad need FSSAI compliance. "
             "Turnover ≤ ₹12 lakh: Registration. > ₹12 lakh: State License. "
             "Online via FoSCoS portal. "
             "Food safety supervisor required for licensed establishments.", False),

            ("2. GHMC Trade License", True),
            ("Trade License from GHMC mandatory. Online application through GHMC Citizen portal. "
             "Categories based on area and nature of business. "
             "Requires: property documents, ID proof, building permission. "
             "Annual renewal. Penalty for late renewal.", False),

            ("3. GHMC Health License", True),
            ("Health/Sanitary License from GHMC required. "
             "Premises inspection by GHMC health officials. "
             "Clean water supply, proper drainage, pest control, waste segregation mandatory. "
             "Medical fitness certificates for all food handlers.", False),

            ("4. Fire NOC", True),
            ("Fire NOC from Telangana State Disaster Response & Fire Services required for: "
             "All commercial buildings above 500 sq ft. "
             "All buildings above 15 meters or 4 floors. "
             "Restaurants with seating capacity above 50. "
             "Requirements: Fire extinguishers — 1 per 1000 sq ft. "
             "Minimum 2 exits for area above 1000 sq ft. "
             "Exit door width: minimum 1.0 meter. "
             "Staircase width: minimum 1.2 meters. "
             "Fire alarm system for buildings above 3 floors. "
             "Sprinkler system for buildings above 5 floors. "
             "Annual renewal of fire safety certificate.", False),

            ("5. Building Permission", True),
            ("Building permission from GHMC/HMDA (Hyderabad Metropolitan Development Authority) required. "
             "Plans must be prepared by a licensed architect. "
             "FAR for commercial: 1.75 to 2.5 depending on road width. "
             "Minimum floor height: 3.0 meters for ground floor, 2.75 meters for upper floors. "
             "Setback: minimum 3 meters front for roads below 12 meters. "
             "HMDA GO Ms No. 168 (Building Rules 2012) applies.", False),

            ("6. Telangana Excise License", True),
            ("Bar License from Telangana State Excise Department for serving liquor. "
             "Minimum area: 1000 sq ft. "
             "Distance restrictions from religious/educational institutions: 100 meters. "
             "Online application through Telangana Excise portal.", False),

            ("7. Police NOC", True),
            ("Police NOC required for restaurants operating after 11 PM. "
             "CCTV surveillance mandatory. "
             "Background verification of owner and staff.", False),

            ("8. Environmental Compliance", True),
            ("NOC from Telangana State Pollution Control Board (TSPCB) for commercial kitchens with significant emissions. "
             "Kitchen exhaust systems with filters required. "
             "Noise limits: 55 dB day, 45 dB night. "
             "Rainwater harvesting mandatory for plots above 200 sq meters.", False),

            ("Key Compliance Checkpoints for Blueprint Review", True),
            ("Floor plan must show kitchen, dining, exits, restrooms, fire safety equipment. "
             "Kitchen must be present and properly ventilated. "
             "Minimum 2 exits for area above 1000 sq ft. "
             "Parking: as per HMDA norms — 1 car space per 50 sq meters. "
             "Rainwater harvesting system on plan for plots > 200 sq meters. "
             "Green building norms apply for large commercial buildings (> 2000 sq meters).", False),
        ],
    },

    "Kolkata": {
        "filename": "Kolkata Restaurant Licensing & Compliance Guide.docx",
        "content": [
            ("Kolkata Restaurant Licensing & Compliance Guide (Official Sources)", True),
            ("Overview", True),
            ("This document covers restaurant licensing in Kolkata, West Bengal. "
             "Businesses must comply with KMC (Kolkata Municipal Corporation), West Bengal state, and central government regulations.", False),

            ("1. FSSAI License / Registration", True),
            ("All food businesses need FSSAI compliance. "
             "Turnover ≤ ₹12 lakh: Registration. > ₹12 lakh: State License. "
             "Online via FoSCoS. "
             "Display of FSSAI license at premises is mandatory.", False),

            ("2. KMC Trade License", True),
            ("Trade License from KMC mandatory. Apply through KMC e-services portal. "
             "Fee depends on area and type of business. "
             "Required: property deed/lease, NOC from building owner, ID proof, PAN. "
             "Annual renewal. Can be done online.", False),

            ("3. KMC Health License", True),
            ("Health License from KMC Health Department required. "
             "Inspection of premises for hygiene compliance. "
             "Clean water supply certificate needed. "
             "Pest control and waste management compliance. "
             "Medical fitness of food handlers.", False),

            ("4. Fire Safety License", True),
            ("Fire Safety License from West Bengal Fire & Emergency Services required for: "
             "Restaurants with area above 500 sq ft. "
             "Buildings above ground floor. "
             "Seating capacity above 50 persons. "
             "Requirements: Fire extinguishers — 1 per 1000 sq ft minimum. "
             "2 exits mandatory for area above 800 sq ft. "
             "Exit width: minimum 1.0 meter. "
             "Staircase width: minimum 1.2 meters. "
             "Fire-resistant kitchen materials. "
             "Emergency exit lighting. "
             "Annual fire safety inspection.", False),

            ("5. Building Plan Sanction", True),
            ("Building plan sanction from KMC required for new construction or change of use. "
             "Plans prepared by registered architect (registered with Council of Architecture). "
             "FAR for commercial in Kolkata: 2.0 to 3.0 depending on road width and zone. "
             "Minimum floor height: 2.75 meters. "
             "Kolkata Municipal Corporation Building Rules apply. "
             "Heritage zone restrictions apply in certain areas (Esplanade, BBD Bagh, etc.).", False),

            ("6. West Bengal Excise License", True),
            ("ON-2 or ON-3 license from WB Excise Department for serving liquor. "
             "Minimum covered area: 750 sq ft for ON-2, 1000 sq ft for ON-3. "
             "Distance from religious/educational: 100 meters. "
             "Serving hours: 10 AM to midnight.", False),

            ("7. Police Permission", True),
            ("NOC from Kolkata Police for restaurants. "
             "Background verification. "
             "CCTV mandatory. "
             "Entertainment license if live music/performance.", False),

            ("8. Environmental NOC", True),
            ("NOC from West Bengal Pollution Control Board (WBPCB) for significant emissions. "
             "Kitchen exhaust compliance. "
             "Noise limits: 55 dB day, 45 dB night. "
             "Rainwater harvesting recommended for buildings above 500 sq meters.", False),

            ("Key Compliance Checkpoints for Blueprint Review", True),
            ("Floor plan must clearly show: kitchen, dining, exits, restrooms, fire equipment. "
             "Kitchen must be present for food establishments. "
             "Minimum 2 exits for area above 800 sq ft. "
             "Heritage zone buildings may have height and facade restrictions. "
             "Parking as per KMC norms. "
             "Proper ventilation in kitchen area (exhaust system on plan).", False),
        ],
    },
}


def create_docx(filename: str, content: list[tuple[str, bool]]):
    """Create a .docx file. Each tuple is (text, is_heading)."""
    doc = Document()
    for text, is_heading in content:
        if is_heading:
            doc.add_heading(text, level=1 if "Guide" in text else 2)
        else:
            doc.add_paragraph(text)
    filepath = os.path.join(DATA_DIR, filename)
    doc.save(filepath)
    print(f"Created: {filepath}")


if __name__ == "__main__":
    for city, info in CITY_DOCS.items():
        create_docx(info["filename"], info["content"])
    print(f"\nDone! Created {len(CITY_DOCS)} city documents in {DATA_DIR}")
