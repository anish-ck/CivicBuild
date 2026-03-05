"""
RAG (Retrieval-Augmented Generation) service.
Loads city-specific licensing documents, chunks them, stores in ChromaDB
with city metadata, and answers user questions using Groq LLM with
retrieved context. Supports multiple cities across India.
"""

import os
import re
from pathlib import Path

import chromadb
from docx import Document
from openai import OpenAI

GROQ_API_KEY = os.getenv("GROQ_API_KEY", "")
GROQ_BASE_URL = "https://api.groq.com/openai/v1"
MODEL = "llama-3.1-8b-instant"

CHROMA_DIR = os.path.join(os.path.dirname(__file__), "..", "chroma_db")
COLLECTION_NAME = "licensing_docs"   # single collection, city in metadata
DATA_DIR = os.path.join(os.path.dirname(__file__), "..", "data")

CHUNK_SIZE = 500       # characters per chunk
CHUNK_OVERLAP = 50     # overlap between chunks
TOP_K = 3              # number of chunks to retrieve

# Supported cities — maps city name variants to canonical key
CITY_ALIASES: dict[str, str] = {
    "chennai": "chennai",
    "mumbai": "mumbai",
    "bombay": "mumbai",
    "bangalore": "bangalore",
    "bengaluru": "bangalore",
    "delhi": "delhi",
    "new delhi": "delhi",
    "hyderabad": "hyderabad",
    "kolkata": "kolkata",
    "calcutta": "kolkata",
}

# Module-level ChromaDB client and collection
_chroma_client = None  # type: chromadb.PersistentClient | None
_collection = None     # type: chromadb.Collection | None


def _get_groq_client() -> OpenAI:
    return OpenAI(api_key=GROQ_API_KEY, base_url=GROQ_BASE_URL)


# ── Document Loading ──────────────────────────────────────────────────

def _load_docx(filepath: str) -> str:
    """Extract all text from a .docx file (paragraphs + tables)."""
    doc = Document(filepath)
    parts = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    return "\n\n".join(parts)


# ── Text Chunking ────────────────────────────────────────────────────

def _chunk_text(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> list[str]:
    """
    Split text into overlapping chunks of approximately `chunk_size` characters.
    Tries to split on sentence boundaries for cleaner chunks.
    """
    # Split into sentences first
    sentences = re.split(r'(?<=[.!?])\s+', text)
    chunks = []
    current_chunk = ""

    for sentence in sentences:
        if len(current_chunk) + len(sentence) + 1 > chunk_size and current_chunk:
            chunks.append(current_chunk.strip())
            # Keep overlap from the end of the current chunk
            overlap_text = current_chunk[-overlap:] if len(current_chunk) > overlap else current_chunk
            current_chunk = overlap_text + " " + sentence
        else:
            current_chunk = (current_chunk + " " + sentence).strip()

    if current_chunk.strip():
        chunks.append(current_chunk.strip())

    return chunks


# ── ChromaDB Initialization ──────────────────────────────────────────

def _detect_city_from_filename(filename: str) -> str:
    """Detect city from document filename."""
    lower = filename.lower()
    for alias, canonical in CITY_ALIASES.items():
        if alias in lower:
            return canonical
    return "general"


def normalize_city(city: str | None) -> str:
    """Normalize a city string to a canonical key. Defaults to 'chennai'."""
    if not city:
        return "chennai"
    return CITY_ALIASES.get(city.strip().lower(), city.strip().lower())


def get_supported_cities() -> list[str]:
    """Return list of supported canonical city names."""
    return sorted(set(CITY_ALIASES.values()))


def initialize():
    """
    Load ALL .docx licensing documents from data/, chunk them,
    and upsert into ChromaDB with city metadata for filtering.
    Uses ChromaDB's default all-MiniLM-L6-v2 embedding model.
    Safe to call multiple times (idempotent).
    """
    global _chroma_client, _collection

    # Create persistent ChromaDB client
    _chroma_client = chromadb.PersistentClient(path=CHROMA_DIR)

    # Get or create the collection (default embedding function)
    _collection = _chroma_client.get_or_create_collection(
        name=COLLECTION_NAME,
        metadata={"hnsw:space": "cosine"},
    )

    # Skip if already populated
    if _collection.count() > 0:
        print(f"[RAG] Collection '{COLLECTION_NAME}' already has {_collection.count()} chunks. Skipping ingestion.")
        return

    # Find all .docx files
    data_path = Path(DATA_DIR)
    if not data_path.exists():
        print("[RAG] WARNING: data/ directory not found. RAG will not work.")
        return

    docx_files = list(data_path.glob("*.docx"))
    if not docx_files:
        print("[RAG] WARNING: No .docx files found in data/ directory.")
        return

    total_chunks = 0
    all_ids = []
    all_docs = []
    all_metas = []

    for docx_path in docx_files:
        city = _detect_city_from_filename(docx_path.name)
        print(f"[RAG] Loading: {docx_path.name} (city={city})")

        full_text = _load_docx(str(docx_path))
        chunks = _chunk_text(full_text)
        print(f"[RAG]   → {len(chunks)} chunks")

        for i, chunk in enumerate(chunks):
            chunk_id = f"{city}_chunk_{total_chunks + i}"
            all_ids.append(chunk_id)
            all_docs.append(chunk)
            all_metas.append({
                "source": docx_path.name,
                "city": city,
                "chunk_index": i,
            })

        total_chunks += len(chunks)

    # Upsert all chunks
    _collection.upsert(ids=all_ids, documents=all_docs, metadatas=all_metas)
    print(f"[RAG] Ingested {total_chunks} total chunks from {len(docx_files)} documents into '{COLLECTION_NAME}'.")


# ── Query / Answer ────────────────────────────────────────────────────

def ask_question(question: str, language_code: str = "en", city: str | None = None) -> str:
    """
    Answer a question using RAG:
    1. Translate question to English if needed (via Sarvam)
    2. Retrieve top-K relevant chunks from ChromaDB (filtered by city)
    3. Send context + question to Groq LLM
    4. Return answer in the user's language
    """
    collection = _collection
    if collection is None or collection.count() == 0:
        return "RAG knowledge base is not initialized. Please ensure the licensing document is loaded."

    # Normalize city
    city_key = normalize_city(city)

    # Step 1: Translate to English if Hindi/Tamil
    query_for_retrieval = question
    if language_code in ("hi-IN", "ta-IN"):
        try:
            from services.speech_service import translate_to_english
            import asyncio
            # translate_to_english is async, run it synchronously
            loop = asyncio.new_event_loop()
            query_for_retrieval = loop.run_until_complete(
                translate_to_english(question, language_code)
            )
            loop.close()
        except Exception:
            query_for_retrieval = question  # Fallback to original

    # Step 2: Retrieve relevant chunks filtered by city
    where_filter = {"city": city_key}
    results = collection.query(
        query_texts=[query_for_retrieval],
        n_results=TOP_K,
        where=where_filter,
    )

    # Fallback: if no results for this city, try without filter
    if not results["documents"] or not results["documents"][0]:
        results = collection.query(
            query_texts=[query_for_retrieval],
            n_results=TOP_K,
        )

    if not results["documents"] or not results["documents"][0]:
        return "Information not available in current knowledge base."

    context_chunks = results["documents"][0]
    context = "\n\n---\n\n".join(context_chunks)

    # Step 3: Determine response language
    lang_map = {
        "hi-IN": "Hindi",
        "ta-IN": "Tamil",
        "en": "English",
    }
    response_language = lang_map.get(language_code, "English")

    # Step 4: Send to Groq LLM
    city_display = city_key.title() if city_key else "the relevant city"
    system_prompt = f"""Answer the question strictly using the provided context about {city_display} regulations.
If the answer is not found in the context, respond exactly with:
"Information not available in current knowledge base."

Respond in {response_language}.
Be concise and helpful."""

    user_message = f"""Context:
{context}

Question: {question}"""

    if not GROQ_API_KEY:
        return "GROQ_API_KEY is not configured."

    try:
        client = _get_groq_client()
        response = client.chat.completions.create(
            model=MODEL,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_message},
            ],
            temperature=0.2,
            max_tokens=1024,
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        return f"Error generating answer: {str(e)}"


# ── Blueprint Compliance Check ────────────────────────────────────────

def check_blueprint_compliance(blueprint_details: dict, city: str | None = None) -> dict:
    """
    Check extracted blueprint details against city-specific licensing rules.
    Uses RAG retrieval (filtered by city) to find relevant regulations, then asks LLM
    to evaluate compliance and return issues/suggestions.

    Returns:
        {
            "compliant": bool,
            "issues": [{"rule": str, "detail": str, "severity": "critical"|"warning"}],
            "suggestions": [str],
            "summary": str,
        }
    """
    collection = _collection
    if collection is None or collection.count() == 0:
        return {
            "compliant": True,
            "issues": [],
            "suggestions": ["RAG knowledge base not loaded — compliance check skipped."],
            "summary": "Could not verify compliance (knowledge base unavailable).",
        }

    city_key = normalize_city(city)

    # Build a description of the blueprint for retrieval
    total_area = blueprint_details.get("total_area", "unknown")
    overall_width = blueprint_details.get("overall_width", "unknown")
    overall_height = blueprint_details.get("overall_height", "unknown")
    floors = blueprint_details.get("floors", "unknown")
    floor_height = blueprint_details.get("floor_height", "unknown")
    seating = blueprint_details.get("seating_capacity", "unknown")
    exits = blueprint_details.get("number_of_exits", "unknown")
    staircases = blueprint_details.get("number_of_staircases", "unknown")
    kitchen = blueprint_details.get("kitchen_present", "unknown")

    # Use multiple retrieval queries to get broader coverage of rules
    queries = [
        "building plan approval fire NOC seating capacity exits requirements",
        "floor plan kitchen dining exits FSSAI restaurant requirements",
        "building construction area height floors commercial restaurant regulations",
    ]

    where_filter = {"city": city_key}
    all_chunks = []
    seen = set()
    for q in queries:
        results = collection.query(query_texts=[q], n_results=TOP_K, where=where_filter)
        if results["documents"] and results["documents"][0]:
            for chunk in results["documents"][0]:
                if chunk not in seen:
                    seen.add(chunk)
                    all_chunks.append(chunk)

    # Fallback: if no city-specific data found, try without filter
    if not all_chunks:
        for q in queries:
            results = collection.query(query_texts=[q], n_results=TOP_K)
            if results["documents"] and results["documents"][0]:
                for chunk in results["documents"][0]:
                    if chunk not in seen:
                        seen.add(chunk)
                        all_chunks.append(chunk)

    context = "\n\n---\n\n".join(all_chunks)

    blueprint_summary = f"""Blueprint Extracted Details:
- Total Area: {total_area}
- Overall Width: {overall_width}
- Overall Height/Depth: {overall_height}
- Number of Floors: {floors}
- Floor Height: {floor_height}
- Seating Capacity: {seating}
- Number of Exits: {exits}
- Number of Staircases: {staircases}
- Kitchen Present: {kitchen}"""

    city_display = city_key.title()
    system_prompt = f"""You are a {city_display} building compliance expert.
Given the blueprint details and {city_display} restaurant licensing regulations, check whether
the building blueprint meets the regulatory requirements.

Evaluate these key compliance areas:
1. Fire NOC: Required if seating > 50 or large built-up area. Must have adequate exits and staircases.
2. Floor Plan: Must show kitchen, dining area, and exits for FSSAI and Fire NOC compliance.
3. Building Plan: Any new construction needs approved building plans and local authority clearance.
4. Kitchen: Must be present for food service businesses (FSSAI requirement).
5. Exits & Safety: Multi-story buildings need adequate staircases and emergency exits.

Respond ONLY in this exact JSON format (no markdown, no code fences):
{{
  "compliant": true or false,
  "issues": [
    {{"rule": "short rule name", "detail": "explanation of the issue", "severity": "critical" or "warning"}}
  ],
  "suggestions": ["actionable suggestion 1", "actionable suggestion 2"],
  "summary": "one paragraph overall assessment"
}}

If all checks pass, set compliant to true and issues to an empty list.
Be practical and specific. Only flag real issues based on the data provided."""

    user_message = f"""{city_display} Restaurant Licensing Regulations:
{context}

{blueprint_summary}

Check this blueprint against the {city_display} regulations and provide your compliance assessment."""

    if not GROQ_API_KEY:
        return {
            "compliant": True,
            "issues": [],
            "suggestions": ["GROQ_API_KEY not configured — compliance check skipped."],
            "summary": "Could not verify compliance (API key missing).",
        }

    try:
        client = _get_groq_client()
        response = client.chat.completions.create(
            model=MODEL,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_message},
            ],
            temperature=0.1,
            max_tokens=1024,
        )
        raw = response.choices[0].message.content.strip()

        # Parse JSON from LLM response
        import json
        # Strip markdown code fences if present
        if raw.startswith("```"):
            raw = raw.split("\n", 1)[1] if "\n" in raw else raw[3:]
            if raw.endswith("```"):
                raw = raw[:-3]
            raw = raw.strip()

        result = json.loads(raw)

        # Ensure required keys
        return {
            "compliant": result.get("compliant", True),
            "issues": result.get("issues", []),
            "suggestions": result.get("suggestions", []),
            "summary": result.get("summary", "Compliance check completed."),
        }
    except json.JSONDecodeError:
        return {
            "compliant": True,
            "issues": [],
            "suggestions": ["Could not parse compliance result — please review manually."],
            "summary": raw if 'raw' in dir() else "Compliance check returned unparseable result.",
        }
    except Exception as e:
        return {
            "compliant": True,
            "issues": [],
            "suggestions": [f"Compliance check error: {str(e)}"],
            "summary": f"Error during compliance check: {str(e)}",
        }
